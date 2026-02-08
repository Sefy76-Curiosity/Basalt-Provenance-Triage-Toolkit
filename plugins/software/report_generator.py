"""
Report Generator Plugin for Basalt Provenance Toolkit
Auto-generate excavation season reports, IAA submissions, etc.

Author: Sefy Levy
License: CC BY-NC-SA 4.0
"""

PLUGIN_INFO = {
      "category": "software",
    "id": "report_generator",
    "name": "Report Generator",
    "description": "Auto-generate excavation season reports, IAA submissions, and provenance summaries",
    "icon": "📄",
    "version": "1.0",
    "requires": ["python-docx"],  # For Word document generation
    "author": "Sefy Levy"
}

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from collections import Counter

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportGeneratorPlugin:
    """Plugin for generating excavation reports"""
    
    def __init__(self, main_app):
        """Initialize plugin"""
        self.app = main_app
        self.window = None
    
    def open_report_generator_window(self):
        """Open the report generator interface"""
        if not HAS_DOCX:
            messagebox.showerror(
                "Missing Dependency",
                "Report Generator requires python-docx:\n\n"
                "Install with:\n"
                "pip3 install python-docx",
                parent=self.app.root
            )
            return
        
        if self.window and self.window.winfo_exists():
            self.window.lift()
            return
        
        self.window = tk.Toplevel(self.app.root)
        self.window.title("Report Generator")
        self.window.geometry("600x480")
        
        # Make window stay on top
        self.window.transient(self.app.root)
        
        self._create_interface()
        
        # Bring to front
        self.window.lift()
        self.window.focus_force()
    
    def _create_interface(self):
        """Create the report generator interface"""
        # Header
        header = tk.Frame(self.window, bg="#9C27B0")
        header.pack(fill=tk.X)
        
        tk.Label(header,
                text="📄 Excavation Report Generator",
                font=("Arial", 16, "bold"),
                bg="#9C27B0", fg="white",
                pady=5).pack()
        
        tk.Label(header,
                text="Auto-generate professional excavation reports in minutes",
                font=("Arial", 10),
                bg="#9C27B0", fg="white",
                pady=5).pack()
        
        # Main content
        content = tk.Frame(self.window, padx=10, pady=8)
        content.pack(fill=tk.BOTH, expand=True)
        
        # Report type selection
        type_frame = tk.LabelFrame(content, text="Report Type",
                                  font=("Arial", 11, "bold"),
                                  padx=8, pady=5)
        type_frame.pack(fill=tk.X, pady=5)
        
        self.report_type = tk.StringVar(value="season_summary")
        
        report_types = [
            ("season_summary", "Season Summary"),
            ("iaa_submission", "IAA Submission"),
            ("provenance_only", "Provenance Only"),
            ("quick_summary", "Quick Summary")
        ]
        
        for value, label in report_types:
            rb = tk.Radiobutton(type_frame, text=label,
                               variable=self.report_type,
                               value=value,
                               font=("Arial", 9))
            rb.pack(anchor=tk.W, pady=1)
        
        # Report details
        details_frame = tk.LabelFrame(content, text="Report Details",
                                     font=("Arial", 11, "bold"),
                                     padx=8, pady=5)
        details_frame.pack(fill=tk.X, pady=5)
        
        # Site name
        tk.Label(details_frame, text="Site Name:",
                font=("Arial", 9)).grid(row=0, column=0, sticky=tk.W, pady=5)
        
        self.site_name = tk.Entry(details_frame, width=40, font=("Arial", 9))
        self.site_name.grid(row=0, column=1, sticky=tk.W, padx=10)
        self.site_name.insert(0, "Tel Hazor")
        
        # Season/Year
        tk.Label(details_frame, text="Season/Year:",
                font=("Arial", 9)).grid(row=1, column=0, sticky=tk.W, pady=5)
        
        self.season = tk.Entry(details_frame, width=40, font=("Arial", 9))
        self.season.grid(row=1, column=1, sticky=tk.W, padx=10)
        self.season.insert(0, str(datetime.now().year))
        
        # Area
        tk.Label(details_frame, text="Area/Square:",
                font=("Arial", 9)).grid(row=2, column=0, sticky=tk.W, pady=5)
        
        self.area = tk.Entry(details_frame, width=40, font=("Arial", 9))
        self.area.grid(row=2, column=1, sticky=tk.W, padx=10)
        self.area.insert(0, "Area G")
        
        # Director
        tk.Label(details_frame, text="Director:",
                font=("Arial", 9)).grid(row=3, column=0, sticky=tk.W, pady=5)
        
        self.director = tk.Entry(details_frame, width=40, font=("Arial", 9))
        self.director.grid(row=3, column=1, sticky=tk.W, padx=10)
        
        # Preview
        preview_frame = tk.LabelFrame(content, text="Report Preview",
                                     font=("Arial", 11, "bold"),
                                     padx=10, pady=5)
        preview_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        preview_text = (
            "Your report will include:\n\n"
            "• Sample counts and classifications\n"
            "• Provenance distribution (Golan, Jordan, Egyptian, etc.)\n"
            "• Wall thickness analysis\n"
            "• Stratigraphic distribution (if locus data available)\n"
            "• Statistical summaries\n"
            "• Recommendations for future work\n\n"
            f"Total samples in database: {len(self.app.samples)}"
        )
        
        tk.Label(preview_frame, text=preview_text,
                font=("Arial", 9),
                justify=tk.LEFT,
                fg="#555").pack(anchor=tk.W)
        
        # Generate button
        tk.Button(content, text="▶ Generate Report",
                 command=self._generate_report,
                 bg="#4CAF50", fg="white",
                 font=("Arial", 12, "bold"),
                 width=25, height=2).pack(pady=5)
    
    def _generate_report(self):
        """Generate the report"""
        if not self.app.samples:
            messagebox.showwarning("No Data",
                                  "No samples to generate report from")
            return
        
        # Get report details
        site = self.site_name.get() or "Unknown Site"
        season = self.season.get() or str(datetime.now().year)
        area = self.area.get() or "Unknown Area"
        director = self.director.get() or ""
        report_type = self.report_type.get()
        
        # Ask where to save
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            initialfile=f"{site.replace(' ', '_')}_Report_{season}.docx"
        )
        
        if not filename:
            return
        
        # Show progress
        if hasattr(self.app, '_show_progress'):
            self.app._show_progress("Generating report...", 10)
        
        try:
            # Generate based on type
            if report_type == "season_summary":
                self._generate_season_summary(filename, site, season, area, director)
            elif report_type == "iaa_submission":
                self._generate_iaa_submission(filename, site, season, area, director)
            elif report_type == "provenance_only":
                self._generate_provenance_report(filename, site, season, area)
            else:  # quick_summary
                self._generate_quick_summary(filename, site, season, area)
            
            if hasattr(self.app, '_hide_progress'):
                self.app._hide_progress()
            
            messagebox.showinfo("Success",
                              f"Report generated successfully!\n\n"
                              f"Saved to:\n{filename}")
            
        except Exception as e:
            if hasattr(self.app, '_hide_progress'):
                self.app._hide_progress()
            messagebox.showerror("Generation Error",
                               f"Failed to generate report:\n\n{e}")
    
    def _generate_season_summary(self, filename, site, season, area, director):
        """Generate comprehensive season summary"""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{site} - Season {season}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(f'Basalt Provenance Analysis Report', 1)
        doc.add_paragraph(f'Area: {area}')
        if director:
            doc.add_paragraph(f'Director: {director}')
        doc.add_paragraph(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph('')
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        
        total = len(self.app.samples)
        classifications = Counter(s.get('Final_Classification', 'Unclassified')
                                 for s in self.app.samples)
        
        summary = (
            f"This report presents the geochemical provenance analysis of {total} basalt "
            f"samples recovered during the {season} excavation season at {site}, {area}. "
            f"Analysis was conducted using trace element geochemistry (Zr, Nb, Ba, Rb, Cr, Ni) "
            f"to determine geological source regions.\n\n"
        )
        
        doc.add_paragraph(summary)
        
        # Sample Overview
        doc.add_heading('Sample Overview', 1)
        doc.add_paragraph(f'Total Samples Analyzed: {total}')
        doc.add_paragraph('')
        
        # Classifications
        doc.add_heading('Provenance Classifications', 2)
        
        for classification, count in classifications.most_common():
            percentage = (count / total) * 100
            doc.add_paragraph(f'• {classification}: {count} samples ({percentage:.1f}%)')
        
        doc.add_paragraph('')
        
        # Wall thickness analysis
        wall_thicknesses = []
        for sample in self.app.samples:
            wt = sample.get('Wall_Thickness_mm', '')
            if wt:
                try:
                    wall_thicknesses.append(float(wt))
                except (ValueError, TypeError):
                    pass
        
        if wall_thicknesses:
            doc.add_heading('Morphological Analysis', 2)
            avg_wt = sum(wall_thicknesses) / len(wall_thicknesses)
            doc.add_paragraph(
                f'Wall thickness data available for {len(wall_thicknesses)} samples. '
                f'Average wall thickness: {avg_wt:.1f} mm '
                f'(range: {min(wall_thicknesses):.1f} - {max(wall_thicknesses):.1f} mm)'
            )
            doc.add_paragraph('')
        
        # Recommendations
        doc.add_heading('Recommendations', 1)
        
        flagged = sum(1 for s in self.app.samples if s.get('Flag_For_Review') == 'YES')
        if flagged > 0:
            doc.add_paragraph(
                f'• {flagged} samples flagged for additional review due to unusual geochemistry'
            )
        
        doc.add_paragraph(
            '• Continue systematic sampling of basalt artifacts in future seasons'
        )
        doc.add_paragraph(
            '• Consider petrographic analysis for ambiguous samples'
        )
        
        # Footer
        doc.add_paragraph('')
        doc.add_paragraph('')
        footer = doc.add_paragraph(
            'Generated by Basalt Provenance Triage Toolkit v10.1'
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(filename)
    
    def _generate_iaa_submission(self, filename, site, season, area, director):
        """Generate IAA authority submission format"""
        doc = Document()
        
        doc.add_heading('ISRAEL ANTIQUITIES AUTHORITY', 0)
        doc.add_heading('Basalt Artifact Analysis Report', 1)
        
        doc.add_paragraph(f'Site: {site}')
        doc.add_paragraph(f'Season: {season}')
        doc.add_paragraph(f'Area: {area}')
        if director:
            doc.add_paragraph(f'Director: {director}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph('')
        
        doc.add_heading('SUMMARY OF FINDINGS', 1)
        doc.add_paragraph(f'Total basalt artifacts analyzed: {len(self.app.samples)}')
        
        # Add classification table
        doc.add_heading('PROVENANCE DETERMINATION', 1)
        classifications = Counter(s.get('Final_Classification', 'Unclassified')
                                 for s in self.app.samples)
        
        for classification, count in classifications.most_common():
            doc.add_paragraph(f'{classification}: {count} artifacts')
        
        doc.add_paragraph('')
        doc.add_paragraph('Analysis conducted using trace element geochemistry.')
        
        doc.save(filename)
    
    def _generate_provenance_report(self, filename, site, season, area):
        """Generate provenance-only report"""
        doc = Document()
        
        doc.add_heading(f'Basalt Provenance Analysis - {site}', 0)
        doc.add_paragraph(f'Season: {season} | Area: {area}')
        doc.add_paragraph('')
        
        classifications = Counter(s.get('Final_Classification', 'Unclassified')
                                 for s in self.app.samples)
        
        for classification, count in classifications.most_common():
            percentage = (count / len(self.app.samples)) * 100
            doc.add_paragraph(f'{classification}: {count} ({percentage:.1f}%)')
        
        doc.save(filename)
    
    def _generate_quick_summary(self, filename, site, season, area):
        """Generate quick summary for field notes"""
        doc = Document()
        
        doc.add_heading(f'{site} - Basalt Summary', 1)
        doc.add_paragraph(f'{season} | {area}')
        doc.add_paragraph('')
        doc.add_paragraph(f'Total samples: {len(self.app.samples)}')
        
        classifications = Counter(s.get('Final_Classification', 'Unclassified')
                                 for s in self.app.samples)
        
        for classification, count in classifications.most_common():
            doc.add_paragraph(f'• {classification}: {count}')
        
        doc.save(filename)
